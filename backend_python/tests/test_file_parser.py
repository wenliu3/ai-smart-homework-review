"""Unit tests for app.core.file_parser 文档文本提取."""
from docx import Document

from app.core import file_parser


class TestExtractTxt:
    def test_reads_txt(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_text("你好 world", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), ".txt") == "你好 world"

    def test_ext_without_leading_dot(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_text("plain", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), "txt") == "plain"

    def test_ext_case_insensitive(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_text("upper", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), ".TXT") == "upper"


class TestExtractDocx:
    def test_reads_docx_paragraphs(self, tmp_path):
        f = tmp_path / "a.docx"
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("second line")
        doc.save(str(f))
        result = file_parser.extract_file_text(str(f), ".docx")
        assert result == "第一段\nsecond line"


class TestFallbacks:
    def test_unknown_extension_returns_empty(self, tmp_path):
        f = tmp_path / "a.md"
        f.write_text("content", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), ".md") == ""

    def test_none_extension_returns_empty(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_text("content", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), None) == ""

    def test_missing_txt_file_swallows_error(self):
        assert file_parser.extract_file_text("/no/such/file.txt", ".txt") == ""

    def test_corrupt_docx_swallows_error(self, tmp_path):
        f = tmp_path / "bad.docx"
        f.write_text("not really a docx", encoding="utf-8")
        assert file_parser.extract_file_text(str(f), ".docx") == ""
