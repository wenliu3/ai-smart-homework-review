"""docx 命中片段标黄 —— 在 docx 文件层面对命中段落加黄色底纹（w:shd）。

对比预览时，后端基于查重命中的公共 N-gram 片段（snippets），遍历 docx 的
正文段落与表格单元格段落（含嵌套表格），命中率达阈值的整段加黄色背景，
生成标黄后的新 docx，再由 LibreOffice 转 PDF 供前端 iframe 预览。

标黄粒度与前端 DOM 级标黄（plagiarism/index.vue 的 applyDomHighlighting）保持一致：
按段落为块，命中 N-gram 比例 >= HIT_RATIO_THRESHOLD 且命中数 >= MIN_HITS 才整段标黄。

注意: 实验报告等内容大量存在于表格中，因此必须覆盖 doc.tables 里每个 cell.paragraphs，
不能只遍历 doc.paragraphs，否则大部分内容不会被标黄。
"""
import os
import shutil
import logging
import subprocess
import tempfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .text import get_char_ngram_set
from .config import PHRASE_NGRAM

logger = logging.getLogger(__name__)

# 与前端 applyDomHighlighting 一致的标黄阈值
HIT_RATIO_THRESHOLD = 0.6
MIN_HITS = 2
HIGHLIGHT_FILL = "FFFF00"  # 黄色底纹


def _set_run_shading(run, fill=HIGHLIGHT_FILL):
    """给 run 加背景底纹（w:shd），已有则先清理再设置，避免重复堆积。"""
    rPr = run._element.get_or_add_rPr()
    for existing in rPr.findall(qn("w:shd")):
        rPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _iter_table_paragraphs(table, seen_cells):
    """递归遍历表格（含嵌套表格）的所有单元格段落。

    合并单元格在 python-docx 中会让 row.cells 返回同一个 cell 的多个包装对象，
    这里把 cell 的 XML 元素本身放入集合去重（lxml element 的 hash 基于 C 节点
    指针，跨代理对象稳定，比 id() 可靠），避免同一段落被重复处理。
    """
    for row in table.rows:
        for cell in row.cells:
            if cell._element in seen_cells:
                continue
            seen_cells.add(cell._element)
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested, seen_cells)


def _iter_all_paragraphs(doc):
    """遍历 doc 的所有段落：正文段落 + 表格单元格段落（含嵌套表格）。

    这是关键修正：原来只遍历 doc.paragraphs 会漏掉表格里的文字，
    而实验报告文档表格占比很高，必须把表格 cell.paragraphs 也纳入。
    """
    for p in doc.paragraphs:
        yield p
    seen_cells = set()
    for table in doc.tables:
        yield from _iter_table_paragraphs(table, seen_cells)


def _paragraph_should_highlight(text, snippet_set, n=PHRASE_NGRAM):
    """判断该段落是否应标黄。

    与前端一致：命中率（命中 n-gram 数 / 该段 n-gram 总数）>= HIT_RATIO_THRESHOLD
    且命中数 >= MIN_HITS 才整段标黄。
    返回 (是否标黄, 命中数, 命中率)。
    """
    if not text or not text.strip() or not snippet_set:
        return False, 0, 0.0
    para_ngrams = get_char_ngram_set(text, n)
    if not para_ngrams:
        return False, 0, 0.0
    hits = sum(1 for g in para_ngrams if g in snippet_set)
    ratio = hits / len(para_ngrams)
    should = ratio >= HIT_RATIO_THRESHOLD and hits >= MIN_HITS
    return should, hits, ratio


def highlight_docx(src_docx_path, snippets, out_docx_path, fill=HIGHLIGHT_FILL):
    """对 src docx 命中片段标黄，另存为 out docx。

    参数:
        src_docx_path: 原始 docx 路径
        snippets: 命中的公共 N-gram 片段列表（与 compare 计算的 snippets 一致）
        out_docx_path: 标黄后输出 docx 路径
        fill: 底纹颜色（默认黄色 FFFF00）

    返回:
        {"highlighted": 标黄段落数, "scanned": 扫描段落数,
         "table_highlighted": 表格内标黄段落数, "details": [{text, hits, ratio, in_table}]}
    """
    snippet_set = set(snippets) if snippets else set()
    try:
        doc = Document(src_docx_path)
    except Exception as e:
        logger.warning("打开 docx 失败: %s: %s", src_docx_path, e)
        raise

    # 先建立正文段落元素集合，用于判断某段落是否属于表格。
    # 用 lxml element 本身入集合（hash 基于 C 节点指针，跨代理稳定），
    # 不能用 id(p._element) —— lxml element 是代理对象，每次访问可能返回不同代理。
    body_elements = set(p._element for p in doc.paragraphs)

    highlighted = 0
    table_highlighted = 0
    scanned = 0
    details = []
    for p in _iter_all_paragraphs(doc):
        scanned += 1
        should, hits, ratio = _paragraph_should_highlight(p.text, snippet_set)
        if not should:
            continue
        in_table = p._element not in body_elements
        for run in p.runs:
            _set_run_shading(run, fill)
        highlighted += 1
        if in_table:
            table_highlighted += 1
        details.append({
            "text": p.text[:80],
            "hits": hits,
            "ratio": round(ratio, 3),
            "in_table": in_table,
        })

    doc.save(out_docx_path)
    logger.info(
        "标黄完成: %s -> %s, 扫描%d段, 标黄%d段(其中表格%d段)",
        src_docx_path, out_docx_path, scanned, highlighted, table_highlighted,
    )
    return {
        "highlighted": highlighted,
        "scanned": scanned,
        "table_highlighted": table_highlighted,
        "details": details,
    }


# ===================== docx → PDF（LibreOffice headless） =====================

def convert_docx_to_pdf(docx_path, output_dir=None, timeout=120):
    """用 LibreOffice headless 把 docx 转成 PDF。

    参数:
        docx_path: docx 文件路径
        output_dir: PDF 输出目录，默认与 docx 同目录
        timeout: 转换超时秒数（soffice 偶尔会卡住）

    返回:
        生成的 PDF 文件绝对路径

    异常:
        任何失败都抛 RuntimeError（不 silent fail），消息含 returncode/stdout/stderr，
        便于在 compare 接口里打印日志定位问题。
    """
    if not os.path.exists(docx_path):
        raise RuntimeError(f"docx 文件不存在: {docx_path}")

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(docx_path))
    os.makedirs(output_dir, exist_ok=True)

    # 每次转换用独立的 UserInstallation 临时目录，避免 soffice 单实例锁导致并发冲突
    profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
    try:
        cmd = [
            "soffice", "--headless", "--norestore", "--nologo", "--nofirststartwizard",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path,
        ]
        logger.info("soffice 转换: %s -> %s", docx_path, output_dir)
        try:
            proc = subprocess.run(
                cmd, capture_output=True, text=True, timeout=timeout, check=False,
            )
        except FileNotFoundError:
            raise RuntimeError(
                "soffice 命令不存在 —— LibreOffice 未安装（Docker 镜像需装 libreoffice-writer）"
            ) from None
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"soffice 转换超时（>{timeout}s）: {docx_path}"
            ) from None

        pdf_path = os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        if proc.returncode != 0 or not os.path.exists(pdf_path):
            raise RuntimeError(
                f"soffice 转换失败 returncode={proc.returncode} | "
                f"stdout={proc.stdout!r} | stderr={proc.stderr!r} | "
                f"pdf存在={os.path.exists(pdf_path)}"
            )
        logger.info("soffice 转换成功: %s", pdf_path)
        return pdf_path
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


# ===================== 标黄 PDF 一体化（标黄 docx → 转 PDF → 缓存） =====================

def make_highlight_cache_key(namespace, self_id, other_id):
    """构造标黄 PDF 的缓存 key。

    key 必须同时包含「被标黄的文档自身」和「对比对方」两方标识 ——
    因为同一份 docx 和不同对象对比时 common_ngrams 不同、标黄结果也不同，
    只用 docx_path 做 key 会读到错误的缓存 PDF。

    namespace: 区分场景，如临时查重的 check_id、作业查重的 'sub'
    self_id:   被标黄文档的标识（临时查重=file序号，作业查重=submission_id）
    other_id:  对比对方的标识

    注意: self_id 与 other_id 有序不交换 —— A的标黄版与B的标黄版是两份不同 PDF。
    """
    return f"hl_{namespace}_{self_id}_{other_id}"


def get_highlighted_pdf(src_docx_path, snippets, cache_key, cache_dir, timeout=120):
    """获取某份 docx 的标黄 PDF（带磁盘缓存）。

    流程:
        1. 缓存命中（PDF 已存在且不旧于源 docx）→ 直接返回
        2. 未命中 → highlight_docx 标黄 → convert_docx_to_pdf 转 PDF → 删中间 docx
        3. 任一步失败 → 清理半成品并抛异常（不 silent fail）

    参数:
        src_docx_path: 原始 docx 路径
        snippets:      命中的公共 N-gram 片段
        cache_key:     缓存 key（务必用 make_highlight_cache_key 构造，含双方 id）
        cache_dir:     缓存目录（标黄 PDF 落盘处）
        timeout:       soffice 转换超时

    返回:
        标黄 PDF 文件路径
    """
    os.makedirs(cache_dir, exist_ok=True)
    pdf_path = os.path.join(cache_dir, f"{cache_key}.pdf")
    tmp_docx = os.path.join(cache_dir, f"{cache_key}.docx")

    src_mtime = os.path.getmtime(src_docx_path)
    # 缓存命中：PDF 存在且不旧于源 docx（源文件被替换则缓存失效）
    if os.path.exists(pdf_path) and os.path.getmtime(pdf_path) >= src_mtime:
        logger.info("标黄PDF缓存命中: %s", pdf_path)
        return pdf_path

    # 未命中：标黄 docx → 转 PDF
    logger.info("生成标黄PDF: src=%s, key=%s", src_docx_path, cache_key)
    highlight_docx(src_docx_path, snippets, tmp_docx)
    try:
        convert_docx_to_pdf(tmp_docx, cache_dir, timeout=timeout)
    except Exception:
        # 转换失败：清理半成品，往上抛（不 silent fail，让接口层记日志）
        for p in (tmp_docx, pdf_path):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass
        raise
    # 转换成功后删除中间标黄 docx，只保留 PDF
    try:
        if os.path.exists(tmp_docx):
            os.remove(tmp_docx)
    except Exception:
        pass
    return pdf_path


def safe_build_highlight_pdf(src_docx_path, snippets, cache_key, cache_dir, timeout=120):
    """生成标黄 PDF 的安全包装：失败时 logger.exception 完整打印堆栈并返回 None。

    - 不 silent fail：异常完整进日志（含 src/key/异常类型）
    - 不让接口崩：返回 None 由调用方把 pdfUrl 置 null，前端可回退
    """
    try:
        return get_highlighted_pdf(src_docx_path, snippets, cache_key, cache_dir, timeout)
    except Exception as e:
        logger.exception(
            "生成标黄PDF失败 src=%s cache_key=%s: %s: %s",
            src_docx_path, cache_key, type(e).__name__, e,
        )
        return None
