"""从 docx / txt / pdf 中提取文本和图片的公共函数。

查重模块（text/image）只负责算相似度，不关心怎么从 docx 里挖数据，
职责更单一。以后要支持 PDF/其他格式，只用改这个文件。
"""
import os
import re
import html
import zipfile
import logging
import xml.etree.ElementTree as ET
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


# ===================== 全文文本提取 =====================

def extract_file_text(file_path: str, ext: str) -> str:
    """从 txt / pdf / docx 文件中提取纯文本"""
    ext = (ext or "").lower()
    if not ext.startswith("."):
        ext = "." + ext
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages)
                return text
        if ext == ".docx":
            try:
                return _extract_docx_text(Document(file_path))
            except Exception:
                text, _ = _extract_docx_via_zip(file_path)
                return text
    except Exception as e:
        logger.warning("提取失败 [%s]: %s: %s", ext, file_path, e)
    return ""


def _extract_docx_text(doc) -> str:
    """从已打开的 Document 对象提取全文文本（段落 + 表格 + 页眉页脚）"""
    parts = []

    # 1. 正文段落
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # 2. 表格文本（实验报告等内容常在表格里）
    def _extract_table(table):
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
                for nested_table in cell.tables:
                    _extract_table(nested_table)
            if row_texts:
                parts.append("\n".join(row_texts))

    for table in doc.tables:
        _extract_table(table)

    # 3. 页眉页脚
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    parts.append(p.text)

    return "\n".join(parts)


# ===================== 图片提取 =====================

def extract_images_from_docx(docx_path: str = None, doc=None) -> list:
    """返回该 docx 内所有内嵌图片的原始二进制内容列表。
    可传入已打开的 Document 对象(doc)以复用，避免同一份文件解析两次。"""
    if doc is None:
        doc = Document(docx_path)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                images.append(rel.target_part.blob)
            except Exception:
                continue
    return images


# ===================== 一次性提取文本+图片（复用 Document 对象） =====================

# ===================== CRC 损坏兜底：用 zipfile 直接读取 docx =====================

# OOXML Word 命名空间
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_P = "{%s}p" % _W_NS
_W_T = "{%s}t" % _W_NS


def _read_zip_member_no_crc(zf: zipfile.ZipFile, name: str) -> bytes:
    """从 ZipFile 中读取条目内容，绕过 CRC-32 校验。

    docx 内某个文件（如 image1.jpeg）的 CRC 校验和不匹配时，标准
    zipfile.read() 会抛出 BadZipFile，导致整个文件解析中断。这里通过将
    ZipExtFile._expected_crc 置 None 来跳过校验，允许读取"损坏"但实际
    内容仍可用的条目。
    """
    try:
        f = zf.open(name)
    except (KeyError, Exception):
        return b""
    try:
        f._expected_crc = None  # 禁用 CRC 校验
        return f.read()
    except Exception:
        return b""
    finally:
        try:
            f.close()
        except Exception:
            pass


def _extract_text_from_ooxml(xml_bytes: bytes) -> str:
    """从 OOXML（word/document.xml、header*.xml、footer*.xml）中提取纯文本。

    遍历 <w:p> 段落，拼接其下所有 <w:t> 文本节点。
    XML 解析失败时回退到正则匹配 <w:t> 内容。
    """
    if not xml_bytes:
        return ""
    try:
        root = ET.fromstring(xml_bytes)
        paragraphs = []
        for p in root.iter(_W_P):
            texts = [t.text for t in p.iter(_W_T) if t.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs)
    except Exception:
        text = xml_bytes.decode("utf-8", errors="ignore")
        snippets = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", text)
        return "".join(snippets) if snippets else ""


def _extract_docx_via_zip(file_path: str) -> tuple:
    """兜底方案：当 python-docx 因 CRC-32 等错误无法打开 docx 时，
    直接用 zipfile 读取文本和图片，绕过 CRC 校验。

    返回: (full_text, images)
      - full_text: 从 word/document.xml + header/footer 提取的全文文本
      - images: word/media/ 下所有图片的原始二进制列表 [bytes, ...]
    """
    full_text_parts = []
    images = []
    try:
        zf = zipfile.ZipFile(file_path, "r")
    except Exception as e:
        logger.warning("zipfile 也无法打开: %s: %s", file_path, e)
        return "", []
    try:
        for info in zf.infolist():
            name = info.filename
            # 文本：主文档 + 页眉页脚
            if name == "word/document.xml" or re.match(r"word/(header|footer)\d+\.xml$", name):
                data = _read_zip_member_no_crc(zf, name)
                if data:
                    text = _extract_text_from_ooxml(data)
                    if text:
                        full_text_parts.append(text)
            # 图片：word/media/ 下的所有文件
            elif name.startswith("word/media/"):
                data = _read_zip_member_no_crc(zf, name)
                if data:
                    images.append(data)
    finally:
        zf.close()
    logger.info("zip 兜底解析完成: %s, 文本%d字符, 图片%d张", file_path, len("\n".join(full_text_parts)), len(images))
    return "\n".join(full_text_parts), images


def extract_all_from_docx(file_path: str) -> tuple:
    """从 docx 中同时提取全文文本和内嵌图片，复用同一个 Document 对象。

    返回: (full_text, images)
      - full_text: 全文文本（段落 + 表格 + 页眉页脚）
      - images: docx 内所有内嵌图片的原始二进制内容列表 [bytes, ...]

    当 python-docx 因 CRC-32 校验失败等错误无法打开时，自动回退到
    zipfile 直接读取（绕过 CRC 校验），确保文本和图片仍可提取。
    """
    try:
        doc = Document(file_path)
        full_text = _extract_docx_text(doc)
        images = extract_images_from_docx(doc=doc)
        return full_text, images
    except Exception as e:
        logger.warning("docx 常规解析失败，尝试 zip 兜底: %s: %s", file_path, e)
        return _extract_docx_via_zip(file_path)


def extract_all_from_file(tmp_path: str, ext: str) -> tuple:
    """通用入口：根据文件扩展名选择提取方式。
    对 .docx 使用 extract_all_from_docx 复用同一个 Document 对象；
    对其他格式用 extract_file_text 提取全文，图片为空列表。
    返回: (full_text, images)
    """
    if ext == ".docx":
        return extract_all_from_docx(tmp_path)
    else:
        full_text = extract_file_text(tmp_path, ext)
        return full_text, []


# ===================== docx → HTML（Word 视图） =====================

def docx_to_html(file_path: str) -> str:
    """用 mammoth 将 docx 转换为保留排版的 HTML（标题/表格/加粗/图片等）。
    mammoth 失败（如 docx 内图片 CRC-32 损坏）时，用 zip 兜底提取纯文本，
    并把图片以 base64 内联追加到末尾（位置无法还原，但至少能显示）。"""
    try:
        import mammoth
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_html(f)
            return result.value
    except Exception as e:
        logger.warning("mammoth 转换失败，回退 zip 兜底: %s: %s", file_path, e)
        text, images = _extract_docx_via_zip(file_path)
        # 按段落拆成 <p>（而非整篇一个 <pre>），前端 applyDomHighlighting 才能按段落标黄命中片段
        parts = []
        for _line in text.split("\n"):
            _line = _line.strip()
            if _line:
                parts.append(f"<p>{html.escape(_line)}</p>")
        # CRC 损坏等场景下，zip 兜底能绕过校验读到图片字节，内联成 base64 显示
        for img_bytes in images:
            try:
                import base64 as _b64
                import io as _io
                from PIL import Image as _PILImage
                im = _PILImage.open(_io.BytesIO(img_bytes))
                fmt = (im.format or "PNG").lower()
                mime = "image/jpeg" if fmt in ("jpg", "jpeg") else f"image/{fmt}"
                b64 = _b64.b64encode(img_bytes).decode("ascii")
                parts.append(f'<img src="data:{mime};base64,{b64}" style="max-width:100%;margin:8px 0" />')
            except Exception:
                continue
        return "\n".join(parts)


def file_to_html(file_path: str, ext: str) -> str:
    """将任意附件文件转为 HTML（docx 用 mammoth，其他用 <pre> 纯文本）"""
    ext = (ext or "").lower()
    if not ext.startswith("."):
        ext = "." + ext
    if ext == ".docx":
        return docx_to_html(file_path)
    text = extract_file_text(file_path, ext)
    return "<pre>" + html.escape(text) + "</pre>" if text else ""
